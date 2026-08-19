from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/Miles/Documents/Free State Curtain Parlour/output")
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "Free State Curtain Parlour Website Agreement.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "111111",
    "muted": "555555",
    "light_fill": "F2F4F7",
    "border": "C9CED6",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="C9CED6", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            set_cell_borders(row.cells[idx])


def finalize_table(table, widths):
    set_table_width(table, widths)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_font(run, size=22, bold=True, color=COLORS["ink"])
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(16)
        r2 = p2.add_run(subtitle)
        set_font(r2, size=11, color=COLORS["muted"])


def add_heading(doc, text, level=1):
    style = "Heading {}".format(level)
    p = doc.add_paragraph(text, style=style)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Detail"
    for cell in hdr:
        set_cell_shading(cell, COLORS["light_fill"])
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, bold=True, color=COLORS["ink"])

    rows = [
        ("Client", "Free State Curtain Parlour"),
        ("Service Provider", "[Your Name / Business Name]"),
        ("Agreement Date", "[Insert Date]"),
        ("Project Type", "Custom multi-page React website and existing Google Business Profile optimisation"),
        ("Timeline", "Two weeks or less from official project start, assuming timely client information, feedback and approvals."),
        ("Revision Rounds", "2 rounds of reasonable revisions"),
        ("Total Project Cost", "R10,768"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run, size=10.5, color=COLORS["ink"])
    finalize_table(table, [2.4, 4.1])


def add_pricing_table(doc):
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Service / Cost"
    hdr[1].text = "Amount"
    for cell in hdr:
        set_cell_shading(cell, COLORS["light_fill"])
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, bold=True, color=COLORS["ink"])

    rows = [
        ("Custom Multi-Page React Website", "R8,500"),
        ("Google Business Profile Optimisation", "R2,000"),
        ("Domain", "R268"),
        ("Total Project Cost", "R10,768"),
    ]
    for idx, (label, amount) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = amount
        if idx == len(rows) - 1:
            for cell in cells:
                set_cell_shading(cell, COLORS["light_fill"])
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if cell == cells[1]:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    set_font(run, size=10.5, bold=(idx == len(rows) - 1), color=COLORS["ink"])
    finalize_table(table, [4.75, 1.75])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_payment_table(doc):
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Payment Stage"
    hdr[1].text = "Amount"
    for cell in hdr:
        set_cell_shading(cell, COLORS["light_fill"])
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, bold=True, color=COLORS["ink"])
    rows = [
        ("Initial 50% service payment", "R5,250"),
        ("Domain cost payable with initial project costs", "R268"),
        ("Initial amount payable to begin", "R5,518"),
        ("Final 50% service payment", "R5,250"),
    ]
    for idx, (label, amount) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = amount
        if idx == 2:
            for cell in cells:
                set_cell_shading(cell, COLORS["light_fill"])
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if cell == cells[1]:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    set_font(run, size=10.5, bold=(idx == 2), color=COLORS["ink"])
    finalize_table(table, [4.75, 1.75])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_signature_table(doc):
    table = doc.add_table(rows=4, cols=2)
    labels = [
        ("Service Provider", "Client"),
        ("Name: ______________________________", "Name: ______________________________"),
        ("Signature: __________________________", "Signature: __________________________"),
        ("Date: _______________________________", "Date: _______________________________"),
    ]
    for row, values in zip(table.rows, labels):
        for cell, text in zip(row.cells, values):
            cell.text = text
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(8 if values == labels[0] else 4)
                for run in p.runs:
                    set_font(run, size=10.5, bold=(values == labels[0]), color=COLORS["ink"])
    finalize_table(table, [3.25, 3.25])


def add_clause_guide(doc):
    add_heading(doc, "Appendix A: Clause Guide", 1)
    add_body(
        doc,
        "This appendix is included for the Service Provider's understanding. It can be removed from the client-facing signing copy if preferred.",
    )
    rows = [
        ("Purpose and scope", "Makes clear what is being provided so the project does not drift into undefined extra work."),
        ("Pre-sale concept distinction", "Protects you from presenting the current homepage concept as a completed production website."),
        ("Exclusions", "Prevents assumptions about e-commerce, ongoing SEO, maintenance, paid ads or other services that are not included."),
        ("Pricing and payment", "Shows exactly what the client is paying for, when payment is due and how the domain cost is handled."),
        ("Timeline", "Sets the two-week-or-less target while accounting for client delays, feedback and content handover."),
        ("Two revision rounds", "Gives the client a fair feedback process while preventing unlimited design changes."),
        ("Client responsibilities", "Makes the client responsible for supplying correct information, access and approvals."),
        ("Google and SEO wording", "Avoids unsupported promises about rankings, traffic, enquiries or sales."),
        ("Ownership", "Transfers the final client-specific work after payment while letting you keep reusable methods and tools."),
        ("Third-party services", "Clarifies that platforms such as Google, hosting providers and domain providers are outside your direct control."),
        ("Cancellation and liability", "Defines what happens if the project stops early and limits exposure to indirect losses."),
    ]
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Clause Area"
    hdr[1].text = "Why It Matters"
    for cell in hdr:
        set_cell_shading(cell, COLORS["light_fill"])
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, bold=True, color=COLORS["ink"])
    for label, reason in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = reason
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run, size=10, color=COLORS["ink"])
    finalize_table(table, [2.2, 4.3])


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Free State Curtain Parlour Website Agreement")
    set_font(r, size=8, color="777777")

    add_title(
        doc,
        "Website Design, Development & Google Business Profile Agreement",
        "Prepared for Free State Curtain Parlour | Prepared: 19 August 2026",
    )
    add_body(
        doc,
        "Important note: This document is a practical commercial template, not legal advice. Before signature, it should be reviewed by a qualified South African legal professional and adjusted to match the final confirmed scope and parties.",
    )
    add_summary_table(doc)

    add_heading(doc, "1. Parties", 1)
    add_body(doc, "This Agreement is entered into between:")
    add_bullets(
        doc,
        [
            "Service Provider: [Your Name / Business Name]",
            "Client: Free State Curtain Parlour",
            "Agreement Date: [Insert Date]",
        ],
    )

    add_heading(doc, "2. Purpose Of This Agreement", 1)
    add_body(
        doc,
        "The Client wishes to appoint the Service Provider to design, develop and launch a custom website, and to optimise the Client's existing Google Business Profile. This Agreement records the scope of work, pricing, payment terms, timeline, responsibilities and important project conditions.",
    )

    add_heading(doc, "3. Project Overview", 1)
    add_body(doc, "The project consists of:")
    add_bullets(
        doc,
        [
            "A custom multi-page React website for Free State Curtain Parlour",
            "Google Business Profile optimisation for the Client's existing profile",
            "Domain configuration",
            "Final testing and launch",
        ],
    )
    add_body(
        doc,
        "The website is intended to represent Free State Curtain Parlour's established showroom, services, history and interior offering online. It is a marketing and lead-generation website, not an e-commerce website.",
    )
    add_body(
        doc,
        "The homepage concept already created is a pre-sale design direction. It demonstrates the proposed visual and strategic direction, but it is not the final production website. Full production development and technical optimisation begin only after the project is approved.",
    )

    add_heading(doc, "4. Website Scope", 1)
    add_body(doc, "The custom website service includes:")
    add_bullets(
        doc,
        [
            "Custom website design",
            "Multi-page React development",
            "Responsive layouts for mobile, tablet, laptop and desktop",
            "Homepage and supporting pages",
            "Integration of supplied/original showroom photography",
            "Image optimisation",
            "Performance optimisation",
            "Technical/on-page SEO foundations",
            "Analytics integration",
            "Contact functionality",
            "Browser and device testing",
            "Domain configuration",
            "Production deployment",
            "Final QA and launch",
        ],
    )
    add_body(doc, "The final page structure may include:")
    add_bullets(
        doc,
        [
            "Home",
            "About Us",
            "Our History",
            "Services",
            "Showroom",
            "Showroom Guide",
            "Contact Us",
        ],
    )
    add_body(
        doc,
        "The exact final structure may be refined during the project based on the Client's content, business needs and approved direction.",
    )

    add_heading(doc, "5. Google Business Profile Scope", 1)
    add_body(
        doc,
        "This service applies to the Client's existing Google Business Profile. The Service Provider is not creating the profile from scratch.",
    )
    add_bullets(
        doc,
        [
            "Auditing the existing profile",
            "Checking and improving business information",
            "Improving the profile description",
            "Reviewing relevant categories and services",
            "Improving how the full service offering is represented",
            "Adding or improving showroom imagery where appropriate",
            "Creating several initial Google Business Profile posts",
            "Connecting the completed website where appropriate",
            "Improving the profile's completeness, accuracy and relevance",
        ],
    )
    add_body(
        doc,
        "The objective is to make the profile more accurate, complete and relevant so Google and potential customers have a clearer understanding of what Free State Curtain Parlour offers. This work does not guarantee rankings, map positions, enquiries or sales.",
    )

    add_heading(doc, "6. Exclusions", 1)
    add_body(doc, "Unless agreed in writing, the following are not included:")
    add_bullets(
        doc,
        [
            "E-commerce functionality",
            "Online payment systems",
            "Product catalogue management",
            "Paid advertising",
            "Ongoing SEO campaigns",
            "Ongoing website maintenance",
            "Ongoing Google Business Profile posting",
            "Professional copywriting beyond agreed website content",
            "Professional photography beyond supplied/original available images",
            "Email hosting setup",
            "Logo design or full brand identity work",
            "Legal policy drafting",
            "Guaranteed Google rankings, traffic, enquiries or sales",
        ],
    )
    add_body(doc, "Additional work outside the agreed scope may be quoted separately.")

    doc.add_page_break()
    add_heading(doc, "7. Pricing", 1)
    add_pricing_table(doc)
    add_body(
        doc,
        "The domain fee is an additional external project cost and is not included inside the R8,500 website service fee.",
    )

    add_heading(doc, "8. Payment Terms", 1)
    add_body(
        doc,
        "The professional service fees are R8,500 for the website and R2,000 for Google Business Profile optimisation, totalling R10,500. These service fees are payable 50% upfront to begin the project and 50% once the completed product is ready.",
    )
    add_payment_table(doc)
    add_body(
        doc,
        "The final service payment is due once the completed product is ready for final approval and before public launch.",
    )

    add_heading(doc, "9. Timeline", 1)
    add_body(
        doc,
        "The project timeline is two weeks or less from the official project start date. This timeline depends on the Client providing required information, access, content, feedback and approvals without significant delay. Client or third-party delays may extend the timeline.",
    )

    add_heading(doc, "10. Client Responsibilities", 1)
    add_body(doc, "The Client is responsible for providing:")
    add_bullets(
        doc,
        [
            "Correct business details",
            "Contact details",
            "Address and trading hours",
            "Website text and content where required",
            "Google Business Profile access where required",
            "Timely feedback and approvals",
            "Images, documents and information needed for the project",
            "Confirmation that supplied content and images may legally be used",
        ],
    )
    add_body(doc, "The Client remains responsible for the accuracy of approved business information.")

    add_heading(doc, "11. Revisions And Changes", 1)
    add_body(
        doc,
        "The project includes 2 rounds of reasonable revisions. A revision means refinement to work already agreed upon, such as text adjustments, image swaps or layout refinements.",
    )
    add_body(
        doc,
        "Major changes to the approved direction, new pages, new features, major restructuring, or work outside the agreed scope may be quoted separately.",
    )

    add_heading(doc, "12. Approval", 1)
    add_body(
        doc,
        "The Client will review the completed website before launch. Once the completed product is ready and approved, the final payment becomes due before launch. Approval may be given in writing by email, message or signed confirmation.",
    )

    add_heading(doc, "13. Ownership", 1)
    add_body(
        doc,
        "Once full payment has been received, the Client owns the final website content and final website design created specifically for the Client.",
    )
    add_body(
        doc,
        "The Service Provider retains ownership of reusable tools, development methods, templates, libraries, workflows, code patterns, processes and pre-existing materials that are not unique to the Client. The Service Provider may display the completed project in a portfolio unless the Client requests otherwise in writing.",
    )

    add_heading(doc, "14. Third-Party Services", 1)
    add_body(
        doc,
        "The project may involve third-party services such as domain providers, hosting platforms, Google Maps, Google Business Profile, analytics tools and external platforms. The Service Provider is not responsible for downtime, policy changes, ranking changes, account restrictions, pricing changes or technical issues caused by third-party platforms.",
    )

    add_heading(doc, "15. SEO And Google Visibility", 1)
    add_body(
        doc,
        "The website will include appropriate technical and on-page SEO foundations such as page titles, metadata, heading structure, semantic page structure, descriptive image text where appropriate and locally relevant business information.",
    )
    add_body(
        doc,
        "The Google Business Profile will be improved for accuracy, completeness and relevance. The Service Provider does not guarantee first-page rankings, Google Maps positions, traffic, enquiries, sales or specific search results.",
    )

    add_heading(doc, "16. Analytics", 1)
    add_body(
        doc,
        "Analytics may be added to measure website activity such as visitors, traffic sources, devices, page views and important contact actions. Analytics data depends on third-party tools, privacy settings and user behaviour, and may not capture every visitor or action.",
    )

    add_heading(doc, "17. Privacy And Personal Information", 1)
    add_body(
        doc,
        "The Service Provider will use personal information and account access only for project purposes. Reasonable steps will be taken to protect access details, business information and personal information. The Client is responsible for ensuring that customer information, images and content supplied to the Service Provider may lawfully be used.",
    )

    add_heading(doc, "18. Cancellation", 1)
    add_body(
        doc,
        "If the Client cancels the project after work has begun, the Client is responsible for work completed up to the cancellation date and any committed external costs, including the domain cost. Any refund or outstanding amount will be calculated based on the work completed and costs already committed at the time of cancellation.",
    )

    add_heading(doc, "19. Limitation Of Liability", 1)
    add_body(
        doc,
        "The Service Provider will take reasonable care in performing the services. The Service Provider is not liable for indirect losses, lost profits, lost sales, Google ranking changes, third-party platform issues, or delays caused by the Client or external providers.",
    )

    add_heading(doc, "20. Final Launch", 1)
    add_body(
        doc,
        "The website will be launched after final review, final payment, required domain/deployment details are confirmed, and final checks are complete.",
    )

    add_heading(doc, "21. Entire Agreement", 1)
    add_body(
        doc,
        "This Agreement represents the agreed project terms between the parties. Any changes to scope, pricing or timeline must be agreed in writing.",
    )

    add_heading(doc, "22. Signatures", 1)
    add_signature_table(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_clause_guide(doc)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
